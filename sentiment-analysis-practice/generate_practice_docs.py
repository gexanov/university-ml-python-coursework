from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "Практика_Малюгин_ЮИ"

STUDENT = "Малюгин Юрий Иванович"
GROUP = "25946"
COURSE = "1"
PRACTICE_DATES = "с 02.02.2026 г. по 27.05.2026 г."
TOPIC = "Кейс «Анализ тональности отзывов для улучшения клиентского сервиса»"
PRACTICE_PLACE = (
    "ФГАОУ ВО «Новосибирский национальный исследовательский государственный университет», "
    "Институт интеллектуальной робототехники, кафедра Интеллектуальных систем теплофизики ИИР, "
    "630090, г. Новосибирск, ул. Пирогова, 1"
)

REPORT_NAME = "Отчет_по_учебной_практике_Малюгин_ЮИ.docx"
ARTICLE_NAME = "Статья_Интерпретируемый_анализ_тональности_Малюгин_ЮИ.docx"
TITLE_NOTE_NAME = "Что_дозаполнить_в_титульнике.md"

FIG_FEATURE_IMPORTANCE = ROOT / "ML" / "LGBM" / "feature_importance_top25.png"
FIG_PERMUTATION_IMPORTANCE = ROOT / "ML" / "LGBM" / "permutation_importance_top20.png"
FIG_CONFUSION_MATRIX = ROOT / "ML" / "LGBM" / "confusion_matrix.png"
FIG_XGB_REFIT_CONFUSION_MATRIX = ROOT / "ML" / "refit_full" / "XGBoost_full_refit_confusion_matrix.png"
FIG_XGB_REFIT_ROC = ROOT / "ML" / "refit_full" / "XGBoost_full_refit_roc_curves_ovr.png"
FIG_XGB_REFIT_PR = ROOT / "ML" / "refit_full" / "XGBoost_full_refit_pr_curves.png"
FIG_ROC = ROOT / "ML" / "LGBM" / "roc_curves_ovr.png"
FIG_PR = ROOT / "ML" / "LGBM" / "pr_curves_ovr.png"
FIG_EDA_OVERVIEW = ROOT / "ML" / "EDA" / "eda_output" / "eda_overview.png"
FIG_PDP_POS_NEG = ROOT / "ML" / "LGBM" / "pdp_ice_pos_neg_top_diff.png"
FIG_PDP_NEG = ROOT / "ML" / "LGBM" / "pdp_ice_num_negations.png"
FIG_PDP_NOT_RECOMMEND = ROOT / "ML" / "LGBM" / "pdp_ice_has_not_recommend.png"
FIG_PDP_NUM_CHARS = ROOT / "ML" / "LGBM" / "pdp_ice_num_chars.png"
FIG_PDP_SENTIMENT_LEXICON = ROOT / "ML" / "LGBM" / "pdp_ice_sentiment_lexicon_diff.png"
FIG_PDP_WORD_SVD_3 = ROOT / "ML" / "LGBM" / "pdp_ice_word_svd_3.png"
FIG_PDP_WORD_SVD_22 = ROOT / "ML" / "LGBM" / "pdp_ice_word_svd_22.png"
FIG_PDP_WORD_SVD_6 = ROOT / "ML" / "LGBM" / "pdp_ice_word_svd_6.png"
FIG_PDP_CHAR_SVD_8 = ROOT / "ML" / "LGBM" / "pdp_ice_char_svd_8.png"
FIG_SHAP_NEG = ROOT / "ML" / "LGBM" / "shap_outputs" / "shap_beeswarm_class_minus1.png"
FIG_SHAP_NEUTRAL = ROOT / "ML" / "LGBM" / "shap_outputs" / "shap_beeswarm_class_0.png"
FIG_SHAP_POS = ROOT / "ML" / "LGBM" / "shap_outputs" / "shap_beeswarm_class_1.png"


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(11)


def set_page_number(paragraph) -> None:
    paragraph.clear()
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style_obj = doc.styles[style_name]
        style_obj.font.name = "Times New Roman"
        style_obj._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    first_line_cm: float = 1.25,
    size_pt: int = 12,
    space_after_pt: int = 0,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = Cm(first_line_cm) if first_line_cm else Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    return p


def add_heading(
    doc: Document,
    text: str,
    level: int = 1,
    *,
    page_break_before: bool = False,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level == 1 else 12)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_caption(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        italic=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
        size_pt=11,
        space_after_pt=6,
    )


def add_table_caption(doc: Document, text: str) -> None:
    p = add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_cm=0,
        size_pt=11,
        space_after_pt=3,
    )
    p.paragraph_format.keep_with_next = True


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right"):
                tag = f"w:{edge}"
                element = tc_borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tc_borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "6")
                element.set(qn("w:space"), "0")
                element.set(qn("w:color"), "000000")


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 15.5) -> None:
    if not path.exists():
        add_paragraph(
            doc,
            f"[Иллюстрация не вставлена автоматически: {path.name}]",
            italic=True,
            first_line_cm=0,
        )
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for idx, item in enumerate(headers):
        set_cell_text(hdr_cells[idx], item)
    for row in rows:
        cells = table.add_row().cells
        for idx, item in enumerate(row):
            cells[idx].text = item
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(11)
    set_table_borders(table)
    return table


def write_report_title_page(doc: Document) -> None:
    add_paragraph(
        doc,
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "РОССИЙСКОЙ ФЕДЕРАЦИИ",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(doc, "", first_line_cm=0)
    add_paragraph(
        doc,
        "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "«НОВОСИБИРСКИЙ НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(doc, "", first_line_cm=0)
    add_paragraph(
        doc,
        "ИНСТИТУТ ИНТЕЛЛЕКТУАЛЬНОЙ РОБОТОТЕХНИКИ",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "Кафедра Интеллектуальных систем теплофизики ИИР",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "Направление подготовки 09.03.03 Прикладная информатика",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "Направленность (профиль) Прикладной искусственный интеллект",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(doc, "", first_line_cm=0)
    add_paragraph(doc, "", first_line_cm=0)
    add_paragraph(
        doc,
        "ОТЧЕТ",
        bold=True,
        size_pt=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "о прохождении учебной практики (ознакомительной практики)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "(указывается наименование практики)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
        size_pt=11,
    )
    add_paragraph(doc, "", first_line_cm=0)
    add_paragraph(
        doc,
        f"Обучающегося {STUDENT} группы № {GROUP} курса {COURSE}",
    )
    add_paragraph(
        doc,
        "(Ф.И.О. полностью)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
        size_pt=11,
    )
    add_paragraph(doc, f"Тема задания: {TOPIC}")
    add_paragraph(doc, "________________________________________________________________________________")
    add_paragraph(doc, "________________________________________________________________________________")
    add_paragraph(doc, f"Место прохождения практики: {PRACTICE_PLACE}")
    add_paragraph(
        doc,
        "(полное наименование организации и структурного подразделения, индекс, адрес)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
        size_pt=11,
    )
    add_paragraph(doc, "________________________________________________________________________________")
    add_paragraph(doc, f"Сроки прохождения практики: {PRACTICE_DATES}")
    add_paragraph(doc, "Руководитель практики__________________________________________________________")
    add_paragraph(doc, "_________________________________________________________/")
    add_paragraph(
        doc,
        "(Ф.И.О. полностью, ученая степень, звание)                     (подпись)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
        size_pt=11,
    )
    add_paragraph(doc, "Оценка по итогам защиты отчета:_________________________________________________")
    add_paragraph(
        doc,
        "(неудовлетворительно, удовлетворительно, хорошо, отлично)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
        size_pt=11,
    )
    add_paragraph(doc, "Отчет заслушан на заседании кафедры              КафИСТИИР")
    add_paragraph(
        doc,
        "(наименование кафедры)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
        size_pt=11,
    )
    add_paragraph(doc, "", first_line_cm=0)
    add_paragraph(doc, "", first_line_cm=0)
    add_paragraph(
        doc,
        "Новосибирск, 2026",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )


def add_contents_line(doc: Document, text: str, page: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(f"{text}\t{page}")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clear_footer_to_single_page_number(section) -> None:
    section.different_first_page_header_footer = True

    for paragraph in list(section.first_page_footer.paragraphs):
        paragraph.clear()

    footer = section.footer
    for paragraph in list(footer.paragraphs[1:]):
        delete_paragraph(paragraph)
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number(p)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_report() -> Document:
    template_path = next(ROOT.glob("отчет*.docx"))
    doc = Document(str(template_path))
    configure_document(doc)
    doc.sections[0].different_first_page_header_footer = True

    title_lines = doc.paragraphs
    title_lines[9].text = "Кафедра интеллектуальных систем теплофизики ИИР"
    title_lines[18].text = f"Обучающегося {STUDENT} группы № {GROUP} курса {COURSE}"
    title_lines[20].text = f"Тема задания: {TOPIC}"
    title_lines[21].text = "______________________________________________________________________________"
    title_lines[22].text = "______________________________________________________________________________"
    title_lines[23].text = "Место прохождения практики: ФГАОУ ВО «Новосибирский национальный исследовательский"
    title_lines[24].text = "государственный университет», Институт интеллектуальной робототехники,"
    title_lines[25].text = "кафедра интеллектуальных систем теплофизики ИИР, 630090, г. Новосибирск, ул. Пирогова, 1"
    title_lines[27].text = f"Сроки прохождения практики: {PRACTICE_DATES}"

    for paragraph in doc.paragraphs[36:48]:
        paragraph.text = ""
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    for paragraph in list(doc.paragraphs[42:48])[::-1]:
        delete_paragraph(paragraph)

    city_idx = next(
        idx for idx, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.strip() == "Новосибирск, 2026"
    )
    for paragraph in list(doc.paragraphs[city_idx + 1:])[::-1]:
        delete_paragraph(paragraph)

    add_page_break(doc)

    add_heading(doc, "СОДЕРЖАНИЕ")
    add_contents_line(doc, "Определения, обозначения и сокращения", "3")
    add_contents_line(doc, "Введение", "4")
    add_contents_line(doc, "1. Краткая характеристика проекта", "6")
    add_contents_line(doc, "2. Выполненные работы в период практики", "8")
    add_contents_line(doc, "3. Полученные результаты и приобретенные навыки", "10")
    add_contents_line(doc, "Заключение", "11")
    add_contents_line(doc, "Список использованных источников и литературы", "12")
    add_contents_line(doc, "Приложение А. Материалы анализа данных и качества модели", "13")
    add_contents_line(doc, "Приложение Б. Материалы интерпретации модели", "19")
    add_page_break(doc)

    add_heading(doc, "ОПРЕДЕЛЕНИЯ, ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ")
    definitions = [
        ("EDA", "exploratory data analysis, разведочный анализ данных: первичная проверка корпуса, распределений, пропусков, дубликатов и базовых статистик текста."),
        ("TF-IDF", "частотное представление текста, учитывающее частоту термина в документе и обратную частоту термина в корпусе."),
        ("SVD", "метод понижения размерности; в проекте использовался для сжатия словесных и символьных TF-IDF признаков."),
        ("Macro-F1", "среднее значение F1-меры по классам без учета их размера; метрика удобна для сравнения качества многоклассовой классификации."),
        ("ROC-AUC OVR", "площадь под ROC-кривой в схеме one-vs-rest для многоклассовой задачи."),
        ("PR-AUC", "площадь под precision-recall кривой; дополнительно показывает качество при разном балансе точности и полноты."),
        ("PDP/ICE", "графики частичной зависимости и индивидуальных условных зависимостей, показывающие, как изменение признака связано с прогнозом модели."),
        ("SHAP", "метод оценки вклада признаков в прогноз модели на основе значений Шепли."),
        ("word_svd_*", "латентные признаки, полученные после сжатия словесного TF-IDF представления методом SVD."),
        ("char_svd_*", "латентные признаки, полученные после сжатия символьного TF-IDF представления методом SVD."),
        ("pos_neg_top_diff", "разность между числом совпадений со словами, характерными для положительного и отрицательного классов."),
        ("num_chars", "длина текста отзыва в символах."),
        ("num_negations", "число отрицаний в тексте отзыва."),
    ]
    for term, meaning in definitions:
        add_paragraph(doc, f"{term} - {meaning}", first_line_cm=0)

    add_heading(doc, "ВВЕДЕНИЕ", page_break_before=True)
    add_paragraph(
        doc,
        "В период учебной практики была выполнена работа, связанная с проектированием и разработкой "
        "программной платформы для автоматического определения тональности русскоязычных отзывов. "
        "Практическая значимость темы определяется тем, что отзывы пользователей являются одним из "
        "наиболее доступных источников сведений о качестве товара, уровне сервиса, типовых проблемах "
        "логистики и соответствии товара ожиданиям покупателя. Для компаний электронной коммерции "
        "оперативная обработка больших массивов отзывов позволяет быстрее выявлять проблемные зоны и "
        "принимать решения на основе данных, а не отдельных субъективных примеров [4, 5]."
    )
    add_paragraph(
        doc,
        "Оформление отчета выполнялось с учетом требований нормоконтроля НГУ и общих правил подготовки "
        "отчетных документов: структуры работы, нумерации страниц, оформления таблиц, рисунков, ссылок "
        "и списка использованных источников [1-3]."
    )
    add_paragraph(
        doc,
        "Целью практики являлось выполнение полного цикла прикладного проекта машинного обучения: "
        "от разведочного анализа датасета и подготовки признаков до обучения модели, интерпретации "
        "ее решений и интеграции в пользовательское приложение. В качестве предметной области был "
        "выбран анализ отзывов о товарах категории «женская одежда и аксессуары» на основе корпуса "
        "RuReviews [10], что позволило работать с содержательно однородным и при этом достаточно "
        "большим корпусом текстов."
    )
    add_paragraph(
        doc,
        "Объектом исследования являются русскоязычные пользовательские отзывы и программная система их "
        "обработки. Предметом исследования выступают методы автоматической классификации тональности "
        "отзывов и интерпретации решений модели машинного обучения."
    )
    add_paragraph(
        doc,
        "Для достижения цели были поставлены следующие задачи: выполнить EDA (exploratory data analysis, "
        "разведочный анализ данных), подготовить признаковое описание текстов, обучить и сравнить "
        "несколько моделей классификации, выбрать наилучший алгоритм по качеству, реализовать backend "
        "и frontend сервиса, а также интерпретировать поведение модели с помощью PDP/ICE "
        "(графиков частичной зависимости и индивидуальных условных зависимостей) и SHAP "
        "(метода оценки вкладов признаков на основе значений Шепли). Отдельное внимание уделялось тому, "
        "чтобы система была пригодна не только для демонстрации, но и для пакетной обработки CSV-файлов "
        "с последующим просмотром, фильтрацией и ручной корректировкой результатов."
    )
    add_paragraph(
        doc,
        "Отдельной задачей являлась проверяемость результата. Поэтому в работе фиксировались не только "
        "итоговые метрики, но и промежуточные артефакты: статистика EDA, параметры разбиения данных, "
        "настроенные гиперпараметры моделей, графики качества, feature importance, permutation importance, "
        "PDP/ICE и SHAP. Такой набор материалов позволяет объяснить, почему выбранная модель была "
        "использована в приложении, и какие признаки действительно влияют на ее решения."
    )
    add_heading(doc, "1. КРАТКАЯ ХАРАКТЕРИСТИКА ПРОЕКТА", page_break_before=True)
    add_paragraph(
        doc,
        "Разработанный проект представляет собой веб-платформу для анализа тональности отзывов. Серверная "
        "часть реализована на FastAPI, клиентская часть — на React. Backend поддерживает одиночный анализ "
        "текста, пакетную обработку CSV-файлов, сохранение результатов на сервере, фильтрацию по источнику "
        "и метке класса, ручную корректировку разметки, расчет метрики macro-F1 и выдачу агрегированных "
        "данных для визуализации. Frontend предоставляет вкладки для всех перечисленных операций и "
        "ориентирован на работу с аналитиком или исследователем."
    )
    add_paragraph(
        doc,
        "В проекте также реализован самостоятельный ML-контур: EDA, подготовка признаков, обучение моделей "
        "CatBoost, XGBoost и LightGBM, расчет метрик качества, построение матриц ошибок, ROC- и PR-кривых, "
        "а также генерация интерпретационных графиков PDP/ICE и SHAP. Такой раздельный подход удобен на "
        "практике: обучение и исследование модели выполняются отдельно, а в приложение поставляются уже "
        "готовые runtime-артефакты."
    )
    add_paragraph(
        doc,
        "Этап EDA включал очистку корпуса и оценку его пригодности для обучения модели. Были удалены "
        "2679 дубликатов и один пустой текст, проверено отсутствие пропусков и утечек между train, "
        "validation и test-выборками, рассчитаны распределения классов и базовые характеристики длины "
        "отзывов. После очистки осталось 87 320 наблюдений, классы сохранили почти равномерное "
        "распределение, а средняя длина текста составила около 133 символов и 20 слов. Эти результаты "
        "использовались при выборе признаков и настройке последующего обучения."
    )
    add_paragraph(
        doc,
        "Признаковое пространство формировалось единым образом для всех сравниваемых моделей. Оно не "
        "подбиралось отдельно под одну модель: логистическая регрессия, CatBoost, XGBoost и LightGBM обучались на одинаковой "
        "подготовленной таблице признаков. В эту таблицу вошли ручные статистические признаки, признаки "
        "лексикона, шаблонные индикаторы, признаки совпадения с частотными словами классов, а также "
        "сжатые компоненты словесного и символьного TF-IDF. Благодаря этому сравнение моделей отражало "
        "именно различие алгоритмов, а не преимущество одной модели из-за отдельного набора признаков [6]."
    )
    add_paragraph(
        doc,
        "Ручные признаки были нужны для того, чтобы модель учитывала простые, но содержательные свойства "
        "отзывов: длину текста, количество слов и предложений, число отрицаний, вопросительных и "
        "восклицательных знаков, наличие повторяющихся символов, соотношение позитивных и негативных "
        "лексических маркеров, а также выражения вроде «не рекомендую» или «не соответствует описанию». "
        "TF-IDF и SVD-компоненты дополняли эти признаки более широким статистическим описанием текста, "
        "которое сложно полностью задать вручную."
    )

    add_heading(doc, "2. ВЫПОЛНЕННЫЕ РАБОТЫ В ПЕРИОД ПРАКТИКИ", page_break_before=True)
    add_paragraph(
        doc,
        "В ходе практики была разработана структура проекта и реализованы основные компоненты сервиса. "
        "Были подготовлены сценарии работы backend и frontend, формат входных данных, логика хранения "
        "результатов пакетного анализа и организация визуализации. В ML-части был реализован общий модуль "
        "формирования признаков, который использовался всеми сравниваемыми моделями. Это важно, потому что "
        "итоговый выбор лучшей по метрике модели был сделан не из-за особого набора признаков, а по результатам "
        "сравнения алгоритмов на одинаковой подготовленной выборке. Дополнительно была проверена линейная "
        "логистическая регрессия как сильный baseline."
    )
    add_paragraph(
        doc,
        "После подготовки признаков были сформированы train, validation и test-таблицы. Для словесного "
        "TF-IDF использовалось ограничение словаря до 8000 признаков, для символьного TF-IDF - до 5000 "
        "признаков. Затем размерность этих представлений была снижена с помощью TruncatedSVD: 200 "
        "компонент для словесного уровня и 150 компонент для символьного уровня. Такой подход позволил "
        "сохранить информацию о лексике и характерных фрагментах слов, но не перегружать бустинговые "
        "модели десятками тысяч разреженных признаков [4, 6]."
    )
    add_paragraph(
        doc,
        "Для каждой бустинговой модели выполнялся подбор гиперпараметров через RandomizedSearchCV с оптимизацией "
        "macro-F1 на кросс-валидации. После подбора модели были дополнительно переобучены на полном train. "
        "Лучшие параметры XGBoost составили: n_estimators = 900, learning_rate = 0.03, max_depth = 8, "
        "min_child_weight = 5, subsample = 1.0, colsample_bytree = 0.7, reg_alpha = 0.0, reg_lambda = 1.0. "
        "Для LightGBM использовались n_estimators = 900, learning_rate = 0.03, num_leaves = 63, max_depth = 10, "
        "min_child_samples = 40, subsample = 0.85, colsample_bytree = 0.7, reg_alpha = 0.5, reg_lambda = 0.0 [8, 9]."
    )
    add_paragraph(
        doc,
        "Был проведен анализ экспериментальных результатов по линейной логистической регрессии и трем бустинговым "
        "моделям: CatBoost, XGBoost и LightGBM. После refit на полном train наилучшее качество среди бустингов "
        "показала XGBoost-модель с macro-F1 = 0.7418, accuracy = 0.7401, ROC-AUC OVR = 0.8932 и PR-AUC = 0.8056. "
        "LightGBM после refit показала близкий результат macro-F1 = 0.7400 и использовалась для детальной "
        "интерпретации PDP/ICE и SHAP."
    )
    add_paragraph(
        doc,
        "Дополнительно оценивались не только итоговые значения метрик, но и устойчивость признаков. "
        "Встроенная важность LightGBM показала высокий вклад SVD-компонент, а permutation importance "
        "проверяла, насколько падает macro-F1 при случайном перемешивании отдельного признака. В таблицу "
        "permutation importance был добавлен нулевой baseline, чтобы отделять признаки с реальным вкладом "
        "от признаков, эффект которых близок к случайному или отсутствует. Наиболее заметное снижение "
        "качества давали латентная словесная SVD-компонента, длина текста, баланс позитивной и негативной "
        "лексики, латентная символьная SVD-компонента и число отрицаний."
    )
    add_paragraph(
        doc,
        "Помимо оценки качества была выполнена содержательная интерпретация модели. Для этого "
        "использовались графики PDP/ICE и SHAP [7]. По ним было установлено, что для данного датасета "
        "ключевую роль играют признаки, связанные с балансом позитивной и негативной лексики, числом "
        "отрицаний, длиной текста, наличием шаблона «не рекомендую», а также скрытыми SVD-компонентами, "
        "кодирующими темы несоответствия описанию, качества ткани, возврата денег, продавца и доставки."
    )
    add_heading(doc, "3. ПОЛУЧЕННЫЕ РЕЗУЛЬТАТЫ И ПРИОБРЕТЕННЫЕ НАВЫКИ", page_break_before=True)
    add_paragraph(
        doc,
        "В результате практики был получен целостный опыт разработки прикладного проекта на стыке "
        "машинного обучения и веб-разработки. Практически были закреплены навыки проектирования "
        "структуры приложения, постановки экспериментов, сравнения моделей, интерпретации результатов, "
        "а также подготовки исследовательских ML-артефактов к использованию в прикладном сервисе."
    )
    add_paragraph(
        doc,
        "Отдельно можно выделить навыки анализа текстовых данных: работа с русскоязычным датасетом "
        "отзывов, оценка качества разметки и баланса классов, использование метрики macro-F1, построение "
        "признакового описания текста и объяснение поведения модели с помощью PDP/ICE и SHAP. Отдельно "
        "был закреплен навык проверки признаков: недостаточно просто получить высокую метрику, нужно "
        "понимать, какие признаки действительно влияют на решение модели и не является ли часть "
        "признакового пространства случайным шумом."
    )
    add_paragraph(
        doc,
        "С практической точки зрения выбранный стек оказался удобным для задачи клиентской аналитики. "
        "Бустинговые модели работают с табличными признаками, быстро выполняют инференс, сохраняются в "
        "виде переносимых артефактов и позволяют применять разные способы интерпретации: встроенную "
        "важность признаков, permutation importance, PDP/ICE и SHAP. Поэтому результат можно использовать "
        "не только как классификатор тональности, но и как инструмент объяснения причин негативных "
        "срабатываний."
    )
    add_paragraph(
        doc,
        "В ходе работы также были получены прикладные навыки интеграции модели в приложение: загрузка "
        "артефактов, проверка входного CSV, пакетный инференс, хранение результатов, фильтрация по меткам, "
        "ручная корректировка спорных предсказаний и расчет агрегированных показателей для интерфейса. "
        "Такая связка ML-экспериментов и сервиса делает проект ближе к реальной задаче клиентского "
        "сервиса, где важны не только метрики, но и удобство последующей работы с результатами."
    )
    add_heading(doc, "ЗАКЛЮЧЕНИЕ", page_break_before=True)
    add_paragraph(
        doc,
        "По итогам учебной практики была разработана платформа анализа тональности русскоязычных отзывов, "
        "объединяющая машинное обучение и веб-интерфейс для прикладной работы с результатами. Поставленные "
        "задачи были решены: выполнен EDA датасета, подготовлены признаки, обучены и сравнены несколько "
        "моделей классификации, включая логистическую регрессию, CatBoost, XGBoost и LightGBM. После refit "
        "лучшей по точечной метрике стала конфигурация XGBoost, а близкая по качеству LightGBM-модель была "
        "использована для интерпретации признаков и подготовки выводов о том, какие языковые сигналы сильнее "
        "всего влияют на решение модели."
    )
    add_paragraph(
        doc,
        "Практика показала, что интерпретируемость является не вспомогательным, а центральным элементом "
        "качественного ML-проекта. Использование PDP и SHAP позволило не только объяснить ответы модели, "
        "но и связать их с конкретными проблемами предметной области: несоответствием товара описанию, "
        "качеством материалов, вопросами доставки и возврата средств. Полученный результат может служить "
        "основой как для дальнейшего развития приложения, так и для подготовки научной публикации."
    )
    add_paragraph(
        doc,
        "Главный итог состоит в том, что модели сравнивались на едином признаковом пространстве, а не за счет "
        "отдельного подбора признаков под одну библиотеку. Это делает вывод о качестве более корректным: "
        "XGBoost после refit показал лучший macro-F1, логистическая регрессия оказалась сильным baseline, "
        "а LightGBM сохранила практическую ценность как близкая по качеству и удобная для интерпретации модель."
    )
    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ И ЛИТЕРАТУРЫ", page_break_before=True)
    references = [
        "1. ГОСТ 7.32-2017. СИБИД. Отчет о научно-исследовательской работе. Структура и правила оформления [Электронный ресурс]. URL: https://docs.cntd.ru/document/1200157208",
        "2. ГОСТ Р 7.0.5-2008. СИБИД. Библиографическая ссылка. Общие требования и правила составления [Электронный ресурс]. URL: https://www.nntu.ru/frontend/web/ngtu/files/org_structura/library/resurvsy/gost_r_7_0_5_2008.pdf",
        "3. НГУ. Нормоконтроль: направление подготовки 09.03.03 Прикладная информатика, направленность «Прикладной искусственный интеллект» : презентация. Новосибирск, 2026.",
        "4. Двойникова А. А., Карпов А. А. Аналитический обзор подходов к распознаванию тональности русскоязычных текстовых данных // Информационно-управляющие системы. 2020. № 4. С. 20-30. DOI: 10.31799/1684-8853-2020-4-20-30.",
        "5. Самигулин Т. Р., Джурабаев А. Э. У. Анализ тональности текста методами машинного обучения // Научный результат. Информационные технологии. 2021. Т. 6. № 1. С. 55-62. DOI: 10.18413/2518-1092-2021-6-1-0-7.",
        "6. Воронцов К. В. Машинное обучение : курс лекций [Электронный ресурс]. URL: https://www.machinelearning.ru/wiki/index.php?title=Машинное_обучение_(курс_лекций,_К.В.Воронцов)",
        "7. Попов Н. В., Шевская Н. В. Методы объяснимого искусственного интеллекта на основе анализа пространства признаков // Компьютерные технологии и системы. 2021. С. 298-300. URL: https://cts.etu.ru/assets/files/2021/cts21/papers/298-300.pdf",
        "8. Яндекс Образование. Градиентный бустинг [Электронный ресурс]. URL: https://education.yandex.ru/handbook/ml/article/gradientnyj-busting",
        "9. Яндекс. CatBoost — библиотека градиентного бустинга на деревьях решений [Электронный ресурс]. URL: https://yandex.ru/dev/catboost/?lang=ru",
        "10. RuReviews: An Automatically Annotated Sentiment Analysis Dataset for Product Reviews in Russian [Электронный ресурс]. URL: https://github.com/sismetanin/rureviews",
    ]
    for item in references:
        add_paragraph(doc, item, first_line_cm=0)
    add_heading(doc, "ПРИЛОЖЕНИЕ А. МАТЕРИАЛЫ АНАЛИЗА ДАННЫХ И КАЧЕСТВА МОДЕЛИ", page_break_before=True)
    add_table_caption(doc, "Таблица А.1 - Сравнение качества моделей на тестовой выборке")
    add_table(
        doc,
        ["Модель", "Accuracy", "Macro-F1", "ROC-AUC OVR", "PR-AUC"],
        [
            ["Logistic Regression", "0.7383", "0.7398", "0.8935", "0.8033"],
            ["CatBoost full refit", "0.7342", "0.7357", "0.8917", "0.8024"],
            ["XGBoost full refit", "0.7401", "0.7418", "0.8932", "0.8056"],
            ["LightGBM full refit", "0.7381", "0.7400", "0.8938", "0.8070"],
        ],
    )
    add_paragraph(
        doc,
        "Для лучшей по точечной метрике модели XGBoost full refit по классам получены следующие F1-показатели: "
        "негативный класс — 0.72, "
        "нейтральный — 0.65, позитивный — 0.85. Наиболее сложным оказался нейтральный класс, что "
        "типично для задач анализа пользовательских отзывов."
    )
    add_paragraph(
        doc,
        "На рисунке А.1 показаны основные результаты EDA: проверка распределения классов, статистика "
        "длины текстов и общая пригодность корпуса к обучению. Этот этап подтвердил, что после удаления "
        "дубликатов данные остаются сбалансированными и могут использоваться без искусственного "
        "перевзвешивания классов."
    )
    add_image(
        doc,
        FIG_EDA_OVERVIEW,
        "Рисунок А.1 - Обзорные результаты EDA по корпусу отзывов",
    )
    add_paragraph(
        doc,
        "На рисунках А.2-А.4 приведены диагностические графики качества XGBoost full refit. Матрица ошибок "
        "показывает, что основная сложность связана с нейтральным классом. ROC- и PR-кривые дополняют "
        "accuracy и macro-F1, так как показывают поведение модели при разных порогах и для каждого класса."
    )
    add_image(
        doc,
        FIG_XGB_REFIT_CONFUSION_MATRIX,
        "Рисунок А.2 - Матрица ошибок XGBoost full refit на тестовой выборке",
    )
    add_image(
        doc,
        FIG_XGB_REFIT_ROC,
        "Рисунок А.3 - ROC-кривые XGBoost full refit в схеме one-vs-rest",
    )
    add_image(
        doc,
        FIG_XGB_REFIT_PR,
        "Рисунок А.4 - Precision-recall кривые XGBoost full refit",
    )
    add_paragraph(
        doc,
        "На рисунках А.5 и А.6 показаны два разных взгляда на важность признаков. Встроенная важность "
        "отражает вклад признаков в структуру деревьев, а permutation importance показывает падение "
        "macro-F1 после перемешивания признака. Совпадение важных признаков в этих подходах усиливает "
        "доверие к интерпретации модели."
    )
    add_image(
        doc,
        FIG_FEATURE_IMPORTANCE,
        "Рисунок А.5 - Встроенная важность признаков LightGBM",
    )
    add_image(
        doc,
        FIG_PERMUTATION_IMPORTANCE,
        "Рисунок А.6 - Permutation importance по снижению macro-F1",
    )
    add_heading(doc, "ПРИЛОЖЕНИЕ Б. МАТЕРИАЛЫ ИНТЕРПРЕТАЦИИ МОДЕЛИ", page_break_before=True)
    add_paragraph(
        doc,
        "Графики PDP/ICE показывают, как модель меняет вероятность классов при изменении конкретного "
        "признака. Рост числа отрицаний и наличие шаблона «не рекомендую» повышают вероятность "
        "негативного класса, а признак pos_neg_top_diff отражает баланс позитивной и негативной лексики."
    )
    add_image(
        doc,
        FIG_PDP_NEG,
        "Рисунок Б.1 - График PDP/ICE для признака num_negations в модели LightGBM",
    )
    add_image(
        doc,
        FIG_PDP_POS_NEG,
        "Рисунок Б.2 - График PDP/ICE для признака pos_neg_top_diff в модели LightGBM",
    )
    add_image(
        doc,
        FIG_PDP_NOT_RECOMMEND,
        "Рисунок Б.3 - График PDP/ICE для признака has_not_recommend в модели LightGBM",
    )
    add_paragraph(
        doc,
        "SHAP-графики показывают распределение вкладов признаков по отдельным объектам. Для отрицательного "
        "класса важны признаки, связанные с негативной лексикой, отрицаниями и латентными компонентами "
        "несоответствия товара описанию. Для положительного класса заметен обратный эффект признаков, "
        "связанных с позитивной лексикой и отсутствием явных жалоб."
    )
    add_image(
        doc,
        FIG_SHAP_NEG,
        "Рисунок Б.4 - SHAP beeswarm для отрицательного класса",
    )
    add_image(
        doc,
        FIG_SHAP_POS,
        "Рисунок Б.5 - SHAP beeswarm для положительного класса",
    )

    clear_footer_to_single_page_number(doc.sections[0])
    return doc


def build_article() -> Document:
    doc = Document()
    configure_document(doc)

    add_paragraph(
        doc,
        "УДК 004.85:004.738.5",
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "ИНТЕРПРЕТИРУЕМАЯ БИЗНЕС-АНАЛИТИКА РУССКОЯЗЫЧНЫХ ОТЗЫВОВ "
        "ЭЛЕКТРОННОЙ КОММЕРЦИИ НА ОСНОВЕ ГРАДИЕНТНОГО БУСТИНГА, SHAP И PDP",
        bold=True,
        size_pt=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "Малюгин Юрий Иванович",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "Новосибирский национальный исследовательский государственный университет, Новосибирск, Россия",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(doc, "", first_line_cm=0)

    add_paragraph(
        doc,
        "Аннотация.",
        bold=True,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "Статья посвящена не только автоматической классификации тональности отзывов, но и задаче "
        "извлечения из модели практически полезной информации о клиентском опыте. На данных RuReviews "
        "по категории «женская одежда и аксессуары» построен воспроизводимый конвейер анализа: "
        "формирование гибридного признакового пространства, сравнение CatBoost, XGBoost и LightGBM, "
        "оценка качества по macro-F1, ROC-AUC и PR-AUC, а также интерпретация лучшего решения с помощью "
        "встроенной важности признаков, permutation importance, SHAP и PDP/ICE. На тестовой выборке "
        "LightGBM показала accuracy 0.7358, macro-F1 0.7378, ROC-AUC OVR 0.8925 и PR-AUC 0.8047. "
        "Главный результат работы состоит в том, что бустинговая модель может использоваться не только "
        "как классификатор тональности, но и как инструмент бизнес-аналитики: интерпретационные методы "
        "выделяют признаки, связанные с несоответствием товара описанию и фотографии, размером, "
        "качеством ткани и швов, отрицаниями, выражением «не рекомендую», доставкой и возвратом денег. "
        "Показано, что SHAP отвечает на вопрос о вкладе признаков в решения модели, permutation importance "
        "проверяет устойчивость этих признаков через падение macro-F1, а PDP/ICE показывает, как изменение "
        "конкретного клиентского сигнала меняет вероятности негативного, нейтрального и позитивного "
        "классов. Такой подход позволяет переходить от автоматической разметки отзывов к объяснимым "
        "выводам о том, что именно важно покупателям."
    )
    add_paragraph(
        doc,
        "Ключевые слова: анализ тональности, электронная коммерция, клиентский опыт, русскоязычные "
        "отзывы, градиентный бустинг, LightGBM, SHAP, PDP, интерпретируемое машинное обучение.",
        italic=True,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "INTERPRETABLE BUSINESS ANALYTICS OF RUSSIAN E-COMMERCE REVIEWS BASED ON GRADIENT BOOSTING, SHAP AND PDP",
        bold=True,
        size_pt=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "Abstract. The paper considers sentiment analysis of Russian product reviews as a business analytics "
        "task rather than only a text classification problem. Using the RuReviews clothing and accessories "
        "subset, we compare CatBoost, XGBoost and LightGBM and interpret the best model with feature "
        "importance, permutation importance, SHAP and PDP/ICE. The results show that gradient boosting "
        "combined with interpretable features can reveal customer pain points related to product mismatch, "
        "size, fabric quality, delivery, refunds and explicit negative recommendation patterns.",
        italic=True,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "Keywords: sentiment analysis, e-commerce, customer experience, Russian reviews, gradient boosting, "
        "LightGBM, SHAP, PDP, interpretable machine learning.",
        italic=True,
        first_line_cm=0,
    )

    add_heading(doc, "1. Введение")
    add_paragraph(
        doc,
        "Отзывы пользователей давно перестали быть только маркетинговым артефактом. Для интернет-магазина "
        "они являются непрерывным источником сведений о качестве товара, соответствии фотографии и описанию, "
        "работе доставки, упаковке, возвратах, спорных ситуациях и готовности покупателя рекомендовать товар. "
        "Поэтому задача анализа тональности практически важна не сама по себе, а как способ быстро выделять "
        "из большого массива текстов повторяющиеся причины удовлетворенности и недовольства клиентов [6, 7]."
    )
    add_paragraph(
        doc,
        "Большая часть прикладных решений по анализу отзывов оценивается через итоговую метрику классификации. "
        "Однако для бизнеса недостаточно получить метку «negative» или «positive»: необходимо понимать, почему "
        "модель отнесла отзыв к определенному классу и какие признаки можно связать с управленческими действиями. "
        "Если негативные срабатывания объясняются словами о возврате средств, несоответствии размера или плохой "
        "ткани, то результат модели превращается в сигнал для товарной аналитики, логистики и клиентского сервиса."
    )
    add_paragraph(
        doc,
        "Современные исследования e-commerce-отзывов развиваются в двух близких направлениях. Первое связано "
        "с аспектно-ориентированным анализом, где заранее выделяются стороны клиентского опыта: доставка, "
        "качество товара, возврат, цена и т. п. [9]. Второе направление использует объяснимое машинное обучение, "
        "в том числе SHAP, чтобы выявлять доменные факторы, влияющие на решения моделей [8]. Настоящая работа "
        "находится между этими подходами: аспекты не размечаются вручную, но извлекаются из интерпретации "
        "бустинговой модели, обученной на русскоязычных отзывах."
    )
    add_paragraph(
        doc,
        "Цель исследования состоит в том, чтобы показать, что градиентный бустинг в сочетании с SHAP, "
        "PDP/ICE и permutation importance может использоваться как практичный инструмент анализа клиентских "
        "отзывов: модель классифицирует тональность, а интерпретационные методы показывают, какие текстовые "
        "и тематические сигналы действительно важны для покупателей. Вклад работы носит прикладной характер: "
        "не предлагается новый алгоритм, но демонстрируется воспроизводимый конвейер, переводящий признаки "
        "модели в содержательные выводы о проблемах категории «одежда и аксессуары»."
    )

    add_heading(doc, "2. Материалы и методы")
    add_paragraph(doc, "2.1. Данные и локальная предобработка", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "Исходный набор данных опирается на датасет RuReviews [4, 5]. Согласно описанию авторов, "
        "полный корпус был автоматически собран из отзывов к товарам категории «Women’s Clothes and "
        "Accessories» на крупной российской e-commerce площадке и первоначально содержал около 821 тыс. "
        "автоматически размеченных отзывов. Для уменьшения дисбаланса классов в открытой версии был "
        "сформирован сбалансированный поднабор из 90 тыс. текстов: по 30 тыс. для отрицательного, "
        "нейтрального и положительного классов."
    )
    add_table_caption(doc, "Таблица 1 — Основные характеристики датасета и локальной предобработки")
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Исходный сбалансированный корпус", "90 000 отзывов"],
            ["Классы", "negative / neutral / positive, по 30 000"],
            ["Удалено дубликатов", "2 679"],
            ["Удалено пустых текстов", "1"],
            ["Итоговый объем после очистки", "87 320"],
            ["Train / validation / test", "61 158 / 13 064 / 13 098"],
            ["Средняя длина текста", "133.24 символа"],
            ["Среднее число слов", "20.50"],
            ["Рекомендуемая max_len", "376"],
        ],
    )
    add_paragraph(
        doc,
        "После удаления 2679 дубликатов и одного пустого текста итоговый объем составил 87 320 отзывов. "
        "Разбиение выполнялось стратифицированно, поэтому в train, validation и test сохранился баланс "
        "трех классов. Для статьи важен не только объем корпуса, но и его предметная однородность: все тексты "
        "относятся к одежде и аксессуарам, поэтому модель может выявлять не абстрактную эмоциональность, а "
        "повторяющиеся доменные сигналы: размер, ткань, цвет, фото, доставка, спор и возврат денег."
    )
    add_paragraph(doc, "2.2. Формирование признаков", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "В проекте использовано гибридное признаковое описание текста. В итоговый вектор вошли 433 признака: "
        "83 ручных статистических и лексических признака, 200 компонент SVD для словесного TF-IDF и 150 "
        "компонент SVD для символьного TF-IDF. Ручные признаки отражают длину текста, число слов, долю "
        "уникальных токенов, число отрицаний, наличие выражения «не рекомендую», долю восклицательных знаков, "
        "словарный баланс позитивной и негативной лексики и совпадения с наиболее частотными словами классов."
    )
    add_paragraph(
        doc,
        "SVD-компоненты применялись не как формальная замена словам, а как способ сжать высокоразмерные "
        "TF-IDF-пространства до плотных признаков, пригодных для бустинга. Словесные компоненты фиксируют "
        "тематические оси на уровне слов и фраз: «не соответствует описанию», «деньги вернули», «качество "
        "не очень», «маломерит». Символьные компоненты полезны для русскоязычных коротких отзывов, потому что "
        "они улавливают основы и фрагменты слов при разных формах, опечатках и смешении вариантов написания: "
        "«достав-», «вернул-», «размер-», «супер», «спор», «открыла». Поэтому словесный SVD дает более "
        "понятные тематические признаки, а символьный SVD добавляет устойчивость к морфологии и неформальному "
        "письму."
    )
    add_paragraph(doc, "2.3. Модели и методы интерпретации", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "Для сравнения были обучены и сохранены три модели: CatBoost, XGBoost и LightGBM. В итоговом "
        "runtime-контуре приложения выбран backend на основе LightGBM [3]. Это решение обосновано двумя "
        "факторами. Во-первых, именно эта модель показала лучший macro-F1 на тестовой выборке. Во-вторых, "
        "ее развертывание удобно с инженерной точки зрения: для инференса не требуется GPU, а рантайм "
        "представляет собой компактный набор сериализованных артефактов, которые могут использоваться "
        "через FastAPI в пакетном режиме."
    )
    add_paragraph(
        doc,
        "Интерпретация строилась на четырех взаимодополняющих инструментах. Встроенная важность LightGBM "
        "показывает, какие признаки часто и результативно использовались в разбиениях деревьев, но может "
        "переоценивать коррелирующие или часто используемые признаки. Permutation importance проверяет "
        "признак жестче: значения признака случайно перемешиваются, после чего измеряется падение macro-F1. "
        "SHAP раскладывает прогноз модели на вклады признаков и позволяет анализировать как глобальную "
        "важность, так и направление влияния на отдельные классы [2]. PDP/ICE, восходящие к идее частичной "
        "зависимости в градиентном бустинге [1], показывают, как меняются вероятности классов при изменении "
        "одного признака при прочих наблюдаемых данных."
    )

    add_heading(doc, "3. Результаты")
    add_paragraph(doc, "3.1. Качество классификации", bold=True, first_line_cm=0)
    add_table_caption(doc, "Таблица 2 — Сравнение качества моделей на тестовой выборке")
    add_table(
        doc,
        ["Модель", "Accuracy", "Precision macro", "Recall macro", "Macro-F1", "ROC-AUC OVR", "PR-AUC"],
        [
            ["CatBoost", "0.7326", "0.7372", "0.7326", "0.7342", "0.8907", "0.8011"],
            ["XGBoost", "0.7351", "0.7399", "0.7351", "0.7369", "0.8922", "0.8039"],
            ["LightGBM", "0.7358", "0.7413", "0.7358", "0.7378", "0.8925", "0.8047"],
        ],
    )
    add_paragraph(
        doc,
        "Различия между моделями оказались не радикальными, что говорит о корректной постановке задачи "
        "и высокой информативности выбранного признакового пространства. Тем не менее LightGBM показала "
        "лучший итоговый macro-F1 = 0.7378. По классам эта модель достигла F1 = 0.72 для отрицательных "
        "отзывов, 0.65 для нейтральных и 0.85 для положительных. Наиболее трудным, как и ожидалось, "
        "оказался нейтральный класс, поскольку он чаще включает краткие, смешанные или малоинформативные "
        "тексты."
    )
    add_paragraph(
        doc,
        "Для прикладного сценария важен не только сам факт лидерства LightGBM, но и характер этого лидерства. "
        "Модель обучалась на том же признаковом пространстве, что CatBoost и XGBoost, поэтому выбор LightGBM "
        "не связан с отдельным подбором признаков под одну библиотеку. В данной задаче LightGBM дала немного "
        "лучший баланс качества, скорости и удобства интерпретации. Такой результат соответствует практической "
        "логике бизнес-аналитики: небольшое преимущество в macro-F1 дополняется возможностью быстро объяснять "
        "решения и пакетно обрабатывать CSV-файлы."
    )
    add_paragraph(doc, "3.2. Сопоставление трех типов важности признаков", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "Встроенная важность LightGBM выделила прежде всего SVD-компоненты: word_svd_3 (1551), word_svd_22 "
        "(1324), char_svd_8 (1244), word_svd_6 (1129), char_svd_4 (1122), char_svd_34 (1066) и word_svd_5 "
        "(1043). Это означает, что деревья активно используют латентные текстовые оси, а не только простые "
        "ручные счетчики. Однако встроенная важность сама по себе не доказывает, что без признака качество "
        "модели заметно ухудшится, поэтому она была сопоставлена с permutation importance и SHAP."
    )
    add_image(
        doc,
        FIG_FEATURE_IMPORTANCE,
        "Рисунок 1. Встроенная важность признаков LightGBM для топ-25 признаков",
    )
    add_paragraph(
        doc,
        "Permutation importance рассчитывалась на валидационной подвыборке с базовым macro-F1 = 0.7279. "
        "Наиболее сильное падение качества дало перемешивание word_svd_3: среднее снижение macro-F1 составило "
        "0.0087. Далее идут num_chars и pos_neg_top_diff с падением 0.0057, char_svd_8 с падением 0.0046, "
        "num_negations с падением 0.0040, char_svd_1 с падением 0.0037, word_svd_22 с падением 0.0030, "
        "char_svd_4 и word_svd_6 с падением около 0.0028, а также exclaims_per_word с падением 0.0028. "
        "Эти значения невелики по абсолютной шкале, но для устойчивой многоклассовой модели на 433 признаках "
        "они показывают, какие признаки реально поддерживают качество."
    )
    add_table_caption(doc, "Таблица 3 — Ключевые признаки по permutation importance")
    add_table(
        doc,
        ["Признак", "Падение macro-F1", "Содержательная интерпретация"],
        [
            ["word_svd_3", "0.0087", "несоответствие описанию, размеру, фото и цвету"],
            ["num_chars", "0.0057", "длина отзыва и степень информативности текста"],
            ["pos_neg_top_diff", "0.0057", "баланс позитивной и негативной лексики"],
            ["char_svd_8", "0.0046", "доставка против благодарственной/позитивной лексики"],
            ["num_negations", "0.0040", "явные отрицания в жалобах и отказах"],
            ["char_svd_1", "0.0037", "возврат денег против размера и качества"],
            ["word_svd_22", "0.0030", "получение заказа, возврат, описание, рекомендация"],
            ["char_svd_4", "0.0028", "продавец/доставка против качества и размера"],
            ["word_svd_6", "0.0028", "качество ткани, синтетика, нитки, плохой материал"],
            ["exclaims_per_word", "0.0028", "эмоциональная насыщенность коротких отзывов"],
        ],
    )
    add_image(
        doc,
        FIG_PERMUTATION_IMPORTANCE,
        "Рисунок 2. Permutation importance по снижению macro-F1",
    )
    add_paragraph(
        doc,
        "SHAP-важность дала близкую, но не полностью совпадающую картину. На первом месте оказался "
        "pos_neg_top_diff со средним абсолютным SHAP-вкладом 0.286. Далее следуют word_svd_3 (0.168), "
        "char_svd_1 (0.134), num_negations (0.124), word_svd_6 (0.100), word_svd_22 (0.099), char_svd_4 "
        "(0.097), char_svd_8 (0.085), word_svd_1 (0.074) и word_svd_5 (0.055). Разница между SHAP и "
        "permutation importance закономерна: SHAP измеряет вклад признака в отдельные прогнозы, а permutation "
        "importance измеряет потерю качества при разрушении признака. Поэтому SHAP лучше показывает, чем модель "
        "пользуется при объяснении решений, а permutation importance — какие признаки труднее заменить "
        "оставшимися признаками."
    )
    add_paragraph(doc, "3.3. Содержательная расшифровка ключевых SVD-компонент", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "SVD-компоненты требуют отдельного пояснения, потому что их названия технические, а смысл возникает "
        "из слов и символьных фрагментов с наибольшими весами на противоположных сторонах компоненты. Знак "
        "SVD-компоненты сам по себе условен: важна не «положительная» или «отрицательная» сторона, а контраст "
        "между группами слов. Именно поэтому компоненту нужно читать как латентную ось, выделяющую повторяющуюся "
        "тему корпуса."
    )
    add_paragraph(
        doc,
        "Компонента word_svd_3 оказалась наиболее важной по permutation importance и первой по встроенной "
        "важности. Ее словесная ось связана с выражениями «не соответствует», «размер не соответствует», "
        "«описанию», «фото», «цвет», «на фото», «картинке» на одной стороне и словами «спасибо», «супер», "
        "«быстрая доставка», «качество хорошее», «отлично» на другой. Поэтому эту компоненту можно трактовать "
        "как ось соответствия ожиданиям: покупатель сравнивает реальный товар с описанием, фотографией, цветом "
        "и размером. Для одежды это один из центральных источников тональности, поскольку товар нельзя полностью "
        "оценить до получения."
    )
    add_paragraph(
        doc,
        "Компонента word_svd_6 описывает качество материала и исполнения. В ее наиболее информативных словах "
        "появляются «качество», «не очень», «плохое», «ужасное», «синтетика», «нитки», «торчат». Это объясняет, "
        "почему она заметна и в SHAP, и в permutation importance: такие слова редко являются нейтральными, а для "
        "категории одежды напрямую отражают товарный дефект. Компонента word_svd_22 смешивает получение заказа, "
        "описание, возврат денег и рекомендацию: «получила», «не получила», «деньги не вернули», «соответствует "
        "описанию», «рекомендую». Ее смысл ближе к операционному опыту покупки, где тональность формируется не "
        "только качеством товара, но и исходом взаимодействия с продавцом."
    )
    add_paragraph(
        doc,
        "Символьные компоненты дополняют словесные. char_svd_1 противопоставляет фрагменты «верну-», «деньг-» "
        "и фрагменты «размер», «качество», поэтому фиксирует ситуации возврата денег на фоне товарных проблем. "
        "char_svd_8 связывает фрагменты «достав-» с благодарственной и позитивной лексикой «спасибо», «отли-», "
        "что отражает роль доставки как самостоятельного фактора удовлетворенности. char_svd_4 содержит "
        "фрагменты, связанные с продавцом и доставкой, а char_svd_34 противопоставляет «супер» и «открыл спор». "
        "Такие компоненты особенно полезны в коротких отзывах, где одно слово или даже устойчивый фрагмент слова "
        "может нести значительную часть смысла."
    )
    add_paragraph(doc, "3.4. PDP/ICE: как признаки меняют вероятности классов", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "PDP/ICE-графики отвечают на другой вопрос, чем SHAP: не «какой вклад внес признак в уже имеющиеся "
        "объекты», а «как в среднем изменится прогноз модели при изменении значения признака». Поэтому они "
        "особенно полезны для ручных признаков, где ось имеет понятную шкалу. Для признака pos_neg_top_diff "
        "рост значения от -4 до 4 уменьшает среднюю вероятность отрицательного класса с 0.3803 до 0.2192 и "
        "увеличивает вероятность положительного класса с 0.2749 до 0.4192. Это почти монотонный эффект: чем "
        "сильнее баланс слов сдвинут к позитивной лексике, тем увереннее модель относит отзыв к положительному "
        "классу."
    )
    add_paragraph(
        doc,
        "Для num_negations увеличение числа отрицаний от 0 до 3 повышает вероятность отрицательного класса "
        "с 0.3075 до 0.3552 и снижает вероятность положительного класса с 0.3539 до 0.3110. Эффект не такой "
        "резкий, как у лексического баланса, потому что отрицание может встречаться и в нейтральных конструкциях, "
        "но в товарных отзывах оно часто связано с жалобой: «не пришел», «не соответствует», «не вернули», "
        "«не рекомендую». Для бинарного паттерна has_not_recommend наличие выражения «не рекомендую» повышает "
        "вероятность отрицательного класса с 0.3295 до 0.3782 и снижает вероятность положительного класса "
        "с 0.3353 до 0.2316. Это показывает, что простые шаблонные признаки остаются ценными даже рядом с "
        "латентными SVD-компонентами."
    )
    add_paragraph(
        doc,
        "Признак num_chars показывает поведение коротких и развернутых отзывов. Для очень коротких текстов "
        "средняя вероятность нейтрального класса выше (0.3641), а вероятность отрицательного ниже (0.2996). "
        "При увеличении длины текста вероятность отрицательного класса стабилизируется около 0.346, а "
        "нейтрального снижается примерно до 0.323. Это согласуется с практической логикой: короткие фразы "
        "часто малоинформативны, тогда как длинные отзывы содержат больше конкретных причин похвалы или жалобы."
    )
    add_image(
        doc,
        FIG_PDP_POS_NEG,
        "Рисунок 3. PDP/ICE для признака pos_neg_top_diff",
    )
    add_image(
        doc,
        FIG_PDP_NEG,
        "Рисунок 4. PDP/ICE для признака num_negations",
    )
    add_image(
        doc,
        FIG_PDP_NOT_RECOMMEND,
        "Рисунок 5. PDP/ICE для признака has_not_recommend",
    )
    add_image(
        doc,
        FIG_PDP_NUM_CHARS,
        "Рисунок 6. PDP/ICE для признака num_chars",
    )
    add_paragraph(
        doc,
        "Для SVD-компонент PDP/ICE читается иначе: ось компоненты не является отдельным словом, а задает "
        "латентный контраст. Для word_svd_3 модель реагирует на движение вдоль оси «несоответствие описанию, "
        "размеру, фото и цвету» против позитивной благодарственной лексики. Для word_svd_6 эффект связан с "
        "осью качества материала: рост выраженности темы «плохое качество», «синтетика», «нитки торчат» "
        "поддерживает негативную интерпретацию. Для char_svd_8 PDP отражает вклад символьной темы доставки, "
        "которая может работать как в позитивном направлении при быстрой доставке, так и в негативном при "
        "долгом ожидании или неполучении заказа."
    )
    add_image(
        doc,
        FIG_PDP_WORD_SVD_3,
        "Рисунок 7. PDP/ICE для латентной компоненты word_svd_3",
    )
    add_image(
        doc,
        FIG_PDP_WORD_SVD_6,
        "Рисунок 8. PDP/ICE для латентной компоненты word_svd_6",
    )
    add_image(
        doc,
        FIG_PDP_CHAR_SVD_8,
        "Рисунок 9. PDP/ICE для латентной компоненты char_svd_8",
    )
    add_paragraph(doc, "3.5. SHAP: вклад признаков по классам", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "SHAP-графики были построены отдельно для отрицательного, нейтрального и положительного классов. "
        "Это важно, потому что один и тот же признак может иметь разный смысл для разных классов. Например, "
        "pos_neg_top_diff при высоких значениях поддерживает положительный класс и одновременно снижает "
        "вероятность отрицательного. Напротив, num_negations и has_not_recommend сильнее связаны с отрицательным "
        "классом. Латентные признаки word_svd_3, word_svd_6 и word_svd_22 показывают, что модель использует не "
        "только отдельные эмоциональные слова, но и устойчивые темы клиентского опыта."
    )
    add_paragraph(
        doc,
        "Отрицательный класс определяется сочетанием явных негативных маркеров и тематических компонент. "
        "Особенно важны признаки несоответствия описанию и фотографии, числа отрицаний, возврата денег, "
        "проблем с размером и качеством ткани. Положительный класс, напротив, поддерживается положительным "
        "лексическим балансом, благодарственной лексикой, словами о соответствии ожиданиям, быстрой доставке "
        "и отсутствии явных жалоб. Нейтральный класс оказывается наиболее сложным: его SHAP-профиль содержит "
        "char_svd_4, exclaims_per_word, word_svd_6 и char_svd_8, что указывает на смешанные и описательные "
        "отзывы, где есть информация о доставке, размере или качестве, но нет сильной эмоциональной полярности."
    )

    add_image(
        doc,
        FIG_SHAP_NEG,
        "Рисунок 10. SHAP beeswarm для отрицательного класса",
    )
    add_image(
        doc,
        FIG_SHAP_NEUTRAL,
        "Рисунок 11. SHAP beeswarm для нейтрального класса",
    )
    add_image(
        doc,
        FIG_SHAP_POS,
        "Рисунок 12. SHAP beeswarm для положительного класса",
    )

    add_heading(doc, "4. Обсуждение")
    add_paragraph(
        doc,
        "Полученные результаты подтверждают основную идею работы: бустинговая модель полезна для бизнеса "
        "не только как быстрый классификатор, но и как основа для объяснимой аналитики отзывов. В отличие "
        "от тяжелых трансформерных моделей, градиентный бустинг работает с табличными признаками, быстро "
        "выполняет инференс, проще разворачивается в backend-сервисе и допускает несколько независимых "
        "способов проверки интерпретируемости. Это не означает, что бустинг всегда лучше трансформеров по "
        "качеству; тезис статьи другой: для прикладной клиентской аналитики важны не только проценты метрики, "
        "но и возможность объяснить, какие проблемы покупателей стоят за предсказаниями."
    )
    add_paragraph(
        doc,
        "На уровне предметной области модель выделяет несколько групп клиентских сигналов. Первая группа — "
        "несоответствие ожиданиям: размер, цвет, фото и описание. Вторая — качество товара: ткань, синтетика, "
        "нитки, швы и общая оценка материала. Третья — логистика и получение заказа: доставка, неполучение, "
        "долгое ожидание. Четвертая — постпокупочный конфликт: спор, возврат денег, фразы «не вернули» и "
        "«не рекомендую». Эти группы хорошо согласуются с фактической природой одежды как товара, где "
        "покупатель особенно чувствителен к посадке, материалу и совпадению реального вида с фотографией."
    )
    add_paragraph(
        doc,
        "С методической точки зрения важна согласованность разных способов интерпретации. Если признак высок "
        "только во встроенной важности, но не влияет на permutation importance и не появляется в SHAP, его "
        "нужно трактовать осторожно. В данной работе ключевые признаки повторяются в нескольких представлениях: "
        "word_svd_3, pos_neg_top_diff, num_negations, char_svd_8, word_svd_6 и word_svd_22 заметны одновременно "
        "в SHAP, permutation importance или встроенной важности. Это повышает доверие к выводу о том, что модель "
        "действительно опирается на содержательные клиентские сигналы, а не на случайный шум."
    )
    add_paragraph(
        doc,
        "Ограничения работы связаны с предметной областью и разметкой. Корпус относится только к одежде и "
        "аксессуарам, поэтому найденные признаки нельзя напрямую переносить на электронику, продукты питания "
        "или услуги. Автоматическая разметка RuReviews также может содержать шум. Кроме того, PDP/ICE для "
        "SVD-компонент требует осторожной интерпретации: изменение латентной оси не равно изменению одного "
        "слова, а отражает движение вдоль группы связанных слов или символьных фрагментов. Тем не менее именно "
        "для анализа одной товарной категории такая специализация является преимуществом: модель показывает "
        "не общую эмоциональность, а конкретные причины, важные для покупателей одежды."
    )

    add_heading(doc, "5. Заключение")
    add_paragraph(
        doc,
        "В работе показано, что градиентный бустинг может быть основой не только для автоматической разметки "
        "тональности русскоязычных отзывов, но и для интерпретируемой бизнес-аналитики клиентского опыта. "
        "На данных категории «женская одежда и аксессуары» лучшей среди трех бустинговых моделей стала "
        "LightGBM с macro-F1 = 0.7378. Однако основной результат состоит не в самой метрике, а в том, что "
        "совместное использование feature importance, permutation importance, SHAP и PDP/ICE позволило "
        "выделить устойчивые группы клиентских сигналов."
    )
    add_paragraph(
        doc,
        "Для бизнеса такие признаки имеют прямую интерпретацию: несоответствие описанию и фотографии, "
        "проблемы размера, качество ткани, доставка, возврат денег, открытие спора и явное выражение "
        "«не рекомендую». SHAP показывает, какие признаки формируют решения модели по классам; permutation "
        "importance подтверждает, какие признаки действительно поддерживают macro-F1; PDP/ICE показывает, "
        "как изменение признака сдвигает вероятности тональности. Поэтому предложенный подход можно "
        "рассматривать как переход от классификации отзывов к объяснимому выявлению клиентских pain points. "
        "Дальнейшая работа может быть связана с проверкой подхода на других товарных категориях, ручной "
        "валидацией выделенных аспектов и сравнением бустингового конвейера с трансформерными моделями."
    )

    add_heading(doc, "Список литературы", page_break_before=True)
    refs = [
        "1. Friedman J. H. Greedy Function Approximation: A Gradient Boosting Machine // Annals of Statistics. 2001. Vol. 29. No. 5. P. 1189-1232.",
        "2. Lundberg S. M., Lee S.-I. A Unified Approach to Interpreting Model Predictions // Advances in Neural Information Processing Systems. 2017. Vol. 30. P. 4765-4774.",
        "3. Ke G., Meng Q., Finley T. et al. LightGBM: A Highly Efficient Gradient Boosting Decision Tree // Advances in Neural Information Processing Systems. 2017. Vol. 30.",
        "4. Smetanin S., Komarov M. Sentiment Analysis of Product Reviews in Russian using Convolutional Neural Networks // 2019 IEEE 21st Conference on Business Informatics (CBI). 2019. P. 482-486. DOI: 10.1109/CBI.2019.00062.",
        "5. RuReviews: An Automatically Annotated Sentiment Analysis Dataset for Product Reviews in Russian [Электронный ресурс]. URL: https://github.com/sismetanin/rureviews",
        "6. Двойникова А. А., Карпов А. А. Аналитический обзор подходов к распознаванию тональности русскоязычных текстовых данных // Информационно-управляющие системы. 2020. № 4. С. 20-30. DOI: 10.31799/1684-8853-2020-4-20-30.",
        "7. Самигулин Т. Р., Джурабаев А. Э. У. Анализ тональности текста методами машинного обучения // Научный результат. Информационные технологии. 2021. Т. 6. № 1. С. 55-62. DOI: 10.18413/2518-1092-2021-6-1-0-7.",
        "8. Mokgwatjane K., Paepae T. An explainable ensemble machine learning approach for multi-domain, multiclass sentiment analysis in Amazon product reviews // Machine Learning with Applications. 2026. Vol. 23. Art. 100825. DOI: 10.1016/j.mlwa.2025.100825.",
        "9. Davoodi L., Mezei J., Heikkilä M. Aspect-based sentiment classification of user reviews to understand customer satisfaction of e-commerce platforms // Electronic Commerce Research. 2026. Vol. 26. P. 1417-1459. DOI: 10.1007/s10660-025-09948-4.",
    ]
    for item in refs:
        add_paragraph(doc, item, first_line_cm=0)

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number(p)
    return doc


def build_article() -> Document:
    doc = Document()
    configure_document(doc)

    add_paragraph(doc, "УДК 004.85:004.738.5", first_line_cm=0)
    add_paragraph(
        doc,
        "ИНТЕРПРЕТИРУЕМАЯ БИЗНЕС-АНАЛИТИКА РУССКОЯЗЫЧНЫХ ОТЗЫВОВ "
        "ЭЛЕКТРОННОЙ КОММЕРЦИИ НА ОСНОВЕ ГРАДИЕНТНОГО БУСТИНГА, SHAP И PDP",
        bold=True,
        size_pt=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(doc, "Малюгин Юрий Иванович", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=0)
    add_paragraph(
        doc,
        "Новосибирский национальный исследовательский государственный университет, Новосибирск, Россия",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(doc, "", first_line_cm=0)

    add_paragraph(doc, "Аннотация.", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "Работа посвящена интерпретируемому анализу тональности русскоязычных отзывов электронной коммерции. "
        "Главная задача состоит не только в классификации отзывов на отрицательные, нейтральные и положительные, "
        "но и в извлечении из модели признаков, которые можно использовать как бизнес-сигналы о клиентском опыте. "
        "На сбалансированном корпусе RuReviews по категории женской одежды и аксессуаров построено признаковое "
        "пространство из 433 признаков: ручные статистические и лексические признаки, 200 SVD-компонент словесного "
        "TF-IDF и 150 SVD-компонент символьного TF-IDF. Для контроля качества сопоставлены логистическая регрессия "
        "на тех же признаках, CatBoost, XGBoost и LightGBM. Логистическая регрессия при сопоставимой схеме обучения "
        "показала macro-F1 = 0.7398, "
        "а лучший после refit бустинг, XGBoost, получил macro-F1 = 0.7418, ROC-AUC OVR = 0.8932 и PR-AUC = 0.8056. "
        "LightGBM после refit показала близкий результат macro-F1 = 0.7400 и использовалась для детальной SHAP/PDP-интерпретации. "
        "Поэтому вклад работы формулируется не как достижение максимальной метрики, а как воспроизводимый способ "
        "превратить модель тональности в объяснимую аналитику отзывов. С помощью feature importance, permutation "
        "importance, SHAP и PDP/ICE показано, что наиболее устойчивые сигналы связаны с несоответствием товара "
        "описанию и фотографии, размером, качеством ткани и швов, доставкой, возвратом денег, отрицаниями и явной "
        "фразой «не рекомендую». Полученные результаты демонстрируют, что градиентный бустинг полезен для бизнеса "
        "как интерпретируемая модель, позволяющая выявлять клиентские pain points в предметной категории.",
    )
    add_paragraph(
        doc,
        "Ключевые слова: анализ тональности, электронная коммерция, клиентский опыт, русскоязычные отзывы, "
        "градиентный бустинг, LightGBM, SHAP, PDP, permutation importance, TF-IDF, SVD.",
        italic=True,
        first_line_cm=0,
    )

    add_page_break(doc)
    add_paragraph(
        doc,
        "INTERPRETABLE BUSINESS ANALYTICS OF RUSSIAN E-COMMERCE REVIEWS BASED ON GRADIENT BOOSTING, SHAP AND PDP",
        bold=True,
        size_pt=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_cm=0,
    )
    add_paragraph(doc, "Abstract.", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "The paper studies sentiment analysis of Russian e-commerce reviews not only as a multiclass text "
        "classification task, but also as a source of interpretable business signals about customer experience. "
        "Using the RuReviews subset for women's clothing and accessories, we construct a 433-dimensional feature "
        "space that combines handcrafted statistical and lexical features, 200 word-level TF-IDF/SVD components "
        "and 150 character-level TF-IDF/SVD components. Logistic regression on the same features, CatBoost, XGBoost "
        "and LightGBM are compared. The logistic baseline reaches macro-F1 = 0.7398 under the same training-subset "
        "scheme, while XGBoost is the best full-refit "
        "boosting model with macro-F1 = 0.7418, ROC-AUC OVR = 0.8932 and PR-AUC = 0.8056. LightGBM reaches a close "
        "macro-F1 = 0.7400 and is used for detailed SHAP/PDP interpretation. Therefore, the contribution "
        "is not claimed as a new state-of-the-art classifier, but as a reproducible interpretable pipeline for "
        "turning review classification into business analytics. Feature importance, permutation importance, SHAP "
        "and PDP/ICE consistently indicate customer pain points related to product-description mismatch, size, "
        "fabric and sewing quality, delivery, refunds, negations and explicit non-recommendation phrases.",
        italic=True,
        first_line_cm=0,
    )
    add_paragraph(
        doc,
        "Keywords: sentiment analysis, e-commerce, customer experience, Russian reviews, gradient boosting, "
        "LightGBM, SHAP, PDP, permutation importance, TF-IDF, SVD.",
        italic=True,
        first_line_cm=0,
    )

    add_heading(doc, "1. Введение")
    add_paragraph(
        doc,
        "Отзывы покупателей в электронной коммерции являются не только текстовой обратной связью, но и массивом "
        "операционных данных о качестве товара, соответствии фотографии и описанию, доставке, упаковке, возвратах "
        "и готовности рекомендовать покупку. Для управления товарной категорией одной метки тональности недостаточно: "
        "бизнесу важно понимать, почему отзыв оказался отрицательным или положительным и какие повторяющиеся причины "
        "стоят за изменением клиентского опыта."
    )
    add_paragraph(
        doc,
        "Существующие исследования анализа тональности [6, 7] часто оценивают качество модели только по итоговой метрике. "
        "Такой подход полезен для сравнения алгоритмов, но хуже отвечает на прикладной вопрос: какие темы в отзывах "
        "реально влияют на решения модели и могут быть преобразованы в управленческие действия. Для категории одежды "
        "такими сигналами могут быть размер, посадка, ткань, швы, соответствие фото, скорость доставки и ситуации "
        "возврата денег."
    )
    add_paragraph(
        doc,
        "В этой работе используется более узкая постановка: не предложить новый алгоритм классификации, а показать, "
        "как градиентный бустинг [1, 3] и методы интерпретации [2] могут работать как практический инструмент бизнес-аналитики "
        "русскоязычных отзывов. Новизна носит прикладной характер: для одной товарной категории строится полный "
        "воспроизводимый конвейер от признаков и подбора моделей до сопоставления нескольких типов важности признаков, "
        "расшифровки SVD-компонент и интерпретации SHAP/PDP в терминах клиентских проблем."
    )
    add_paragraph(
        doc,
        "Цель исследования - проверить, какие текстовые и латентные признаки оказываются значимыми для определения "
        "тональности отзывов и как эти признаки можно интерпретировать для клиентского сервиса. Для этого решаются "
        "четыре задачи: сформировать гибридное признаковое пространство, сравнить линейный baseline и три модели "
        "градиентного бустинга, проверить устойчивость признаков через permutation importance и объяснить поведение "
        "интерпретируемой бустинговой модели с помощью SHAP и PDP/ICE [2]."
    )

    add_heading(doc, "2. Материалы и методы")
    add_paragraph(doc, "2.1. Данные и контроль качества корпуса", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "Исходный корпус основан на RuReviews [4, 5] - датасете русскоязычных отзывов о товарах, автоматически "
        "собранном и размеченном по тональности. В проекте использован сбалансированный поднабор категории "
        "«женская одежда и аксессуары»: по 30 000 отзывов для отрицательного, нейтрального и положительного классов. "
        "На этапе EDA отдельно проверялись пустые тексты, дубликаты, распределение длин и возможные пересечения "
        "между выборками. Для сопоставимости моделей в итоговых экспериментах использовался фиксированный "
        "стратифицированный split 72 000 / 18 000, где в тестовой части содержится по 6 000 отзывов каждого класса."
    )
    add_table_caption(doc, "Таблица 1 - Характеристики данных и экспериментального split")
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Предметная область", "женская одежда и аксессуары"],
            ["Классы", "negative / neutral / positive"],
            ["Исходный сбалансированный корпус", "90 000 отзывов"],
            ["Train / test", "72 000 / 18 000"],
            ["Поддержка классов в test", "по 6 000 отзывов"],
            ["Средняя длина текста по EDA", "133.24 символа"],
            ["Среднее число слов по EDA", "20.50"],
            ["Зафиксированный random_state", "42"],
        ],
    )

    add_paragraph(doc, "2.2. Формирование признакового пространства", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "Признаки строились так, чтобы модель видела не только отдельные слова, но и структуру короткого отзыва: "
        "длину текста, пунктуацию, отрицания, эмоциональные маркеры, шаблонные фразы, лексический баланс и устойчивые "
        "словесные или символьные сочетания. Всего использовались 433 признака: 83 ручных признака, 200 компонент "
        "словесного TF-IDF/SVD и 150 компонент символьного TF-IDF/SVD."
    )
    add_table_caption(doc, "Таблица 2 - Группы признаков")
    add_table(
        doc,
        ["Группа признаков", "Количество", "Содержательный смысл"],
        [
            [
                "Статистические признаки",
                "часть из 83",
                "длина текста, число слов, знаки препинания, заглавные буквы, повторы символов",
            ],
            [
                "Лексические и шаблонные признаки",
                "часть из 83",
                "отрицания, интенсификаторы, слова доставки, качества, размера, возврата, фразы «не рекомендую»",
            ],
            [
                "SVD-компоненты словесного TF-IDF",
                "200",
                "латентные темы на уровне слов и фраз: размер, качество, доставка, возврат, соответствие описанию",
            ],
            [
                "SVD-компоненты символьного TF-IDF",
                "150",
                "фрагменты слов, морфологические варианты, опечатки и короткие устойчивые формы",
            ],
        ],
    )
    add_table_caption(doc, "Таблица 3 - Параметры TF-IDF и SVD")
    add_table(
        doc,
        ["Этап", "Параметры"],
        [
            [
                "Word TF-IDF",
                "lowercase=True; ngram_range=(1, 3); min_df=7; max_df=0.95; max_features=8000; sublinear_tf=True",
            ],
            [
                "Char TF-IDF",
                "analyzer=char_wb; ngram_range=(3, 7); min_df=7; max_df=0.95; max_features=5000; sublinear_tf=True",
            ],
            [
                "Word SVD",
                "TruncatedSVD(n_components=200, random_state=42), fit только на train",
            ],
            [
                "Char SVD",
                "TruncatedSVD(n_components=150, random_state=42), fit только на train",
            ],
        ],
    )
    add_paragraph(
        doc,
        "SVD-компоненты интерпретировались не как отдельные слова, а как латентные оси. Их смысл оценивался по словам "
        "и n-граммам с наибольшими по модулю весами в компоненте. Поэтому выводы по SVD в статье сформулированы "
        "осторожно: компонент word_svd_3 не означает одно конкретное слово, а указывает на группу совместно "
        "встречающихся сигналов, связанных с несоответствием ожиданиям."
    )

    add_paragraph(doc, "2.3. Модели, подбор параметров и интерпретация", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "В качестве контрольной модели обучалась логистическая регрессия на тех же 433 признаках. Для сопоставимости "
        "с сохраненными бустинговыми моделями финальная контрольная оценка считалась на той же схеме с внутренним "
        "validation-разделением train. Для основных "
        "экспериментов использовались CatBoost, XGBoost и LightGBM [1, 3]. Гиперпараметры бустингов подбирались через "
        "RandomizedSearchCV: 20 случайных комбинаций, 3-fold cross-validation, целевая метрика macro-F1, random_state=42. "
        "Такой подбор нужен не для доказательства абсолютного превосходства одного алгоритма, а для честного сравнения "
        "трех бустингов в одинаковом признаковом пространстве."
    )
    add_table_caption(doc, "Таблица 4 - Итоговые настройки моделей")
    add_table(
        doc,
        ["Модель", "Ключевые настройки"],
        [
            ["Logistic Regression", "StandardScaler; solver=lbfgs; max_iter=1000; random_state=42"],
            ["CatBoost", "depth=6; iterations=1200; learning_rate=0.07; l2_leaf_reg=1"],
            [
                "XGBoost",
                "n_estimators=900; learning_rate=0.03; max_depth=8; min_child_weight=5; subsample=1.0; colsample_bytree=0.7",
            ],
            [
                "LightGBM",
                "n_estimators=900; learning_rate=0.03; num_leaves=63; max_depth=10; min_child_samples=40; subsample=0.85; colsample_bytree=0.7; reg_alpha=0.5; reg_lambda=0.0",
            ],
        ],
    )
    add_paragraph(
        doc,
        "Интерпретация проводилась несколькими независимыми способами. Встроенная feature importance показывает, "
        "какие признаки часто используются деревьями. Permutation importance проверяет, насколько падает macro-F1 "
        "после случайного перемешивания одного признака; для контроля добавлен случайный baseline-признак. SHAP "
        "показывает локальный вклад признаков в решение модели по каждому классу на выборке из 300 объектов. PDP/ICE "
        "показывает, как изменение выбранного признака связано с вероятностями классов на выборке из 3000 объектов. "
        "Различие методов принципиально: SHAP отвечает «что повлияло на конкретное предсказание», PDP - «как меняется "
        "средняя реакция модели при изменении признака», а permutation importance - «насколько признак нужен для качества»."
    )

    add_heading(doc, "3. Результаты")
    add_paragraph(doc, "3.1. Качество классификации и устойчивость выводов", bold=True, first_line_cm=0)
    add_table_caption(doc, "Таблица 5 - Качество моделей на тестовой выборке")
    add_table(
        doc,
        ["Модель", "Accuracy", "Precision macro", "Recall macro", "Macro-F1", "95% CI Macro-F1", "ROC-AUC OVR", "PR-AUC"],
        [
            ["Logistic Regression", "0.7383", "0.7426", "0.7383", "0.7398", "[0.7332; 0.7463]", "0.8935", "0.8033"],
            ["CatBoost full refit", "0.7342", "0.7385", "0.7342", "0.7357", "[0.7298; 0.7420]", "0.8917", "0.8024"],
            ["XGBoost full refit", "0.7401", "0.7450", "0.7401", "0.7418", "[0.7352; 0.7477]", "0.8932", "0.8056"],
            ["LightGBM full refit", "0.7381", "0.7433", "0.7381", "0.7400", "[0.7338; 0.7462]", "0.8938", "0.8070"],
        ],
    )
    add_paragraph(
        doc,
        "Сравнение показывает две важные вещи. Во-первых, refit на полном train действительно улучшил бустинги: "
        "XGBoost вырос до macro-F1 = 0.7418, а LightGBM - до 0.7400. Поэтому после refit лучший бустинг уже "
        "немного превосходит логистическую регрессию с macro-F1 = 0.7398. Во-вторых, bootstrap-интервалы моделей "
        "все равно пересекаются, поэтому разницу нельзя трактовать как сильное статистически доказанное превосходство. "
        "Практический вывод остается более осторожным: бустинг дает качество на уровне сильного линейного baseline, "
        "но дополнительно поддерживает нелинейные зависимости и удобную интерпретацию через деревья, SHAP и PDP/ICE."
    )
    add_paragraph(
        doc,
        "По отчету классификации XGBoost full refit хуже всего распознает нейтральный класс: F1 = 0.65 против 0.72 для отрицательного "
        "и 0.85 для положительного. Это ожидаемо для отзывов: нейтральные тексты часто содержат фактическое описание "
        "доставки, размера или качества без сильной эмоциональной оценки, из-за чего они оказываются ближе и к слабому "
        "позитиву, и к слабому негативу."
    )
    add_image(doc, FIG_XGB_REFIT_CONFUSION_MATRIX, "Рисунок 1. Матрица ошибок XGBoost full refit", width_cm=13.5)

    add_page_break(doc)
    add_paragraph(doc, "3.2. Важность признаков и permutation importance", bold=True, first_line_cm=0)
    add_image(doc, FIG_FEATURE_IMPORTANCE, "Рисунок 2. Встроенная важность признаков LightGBM", width_cm=14.5)
    add_image(doc, FIG_PERMUTATION_IMPORTANCE, "Рисунок 3. Permutation importance по снижению macro-F1", width_cm=14.5)
    add_table_caption(doc, "Таблица 6 - Ключевые признаки по permutation importance")
    add_table(
        doc,
        ["Признак", "Падение macro-F1", "Содержательная интерпретация"],
        [
            ["word_svd_3", "0.0087", "несоответствие описанию, размеру, фото и цвету"],
            ["num_chars", "0.0057", "длина и информативность отзыва"],
            ["pos_neg_top_diff", "0.0057", "баланс позитивных и негативных маркеров"],
            ["char_svd_8", "0.0046", "доставка и противопоставление позитивной лексике"],
            ["num_negations", "0.0040", "отрицания в жалобах и отказах"],
            ["char_svd_1", "0.0037", "возврат денег и проблемы с качеством"],
            ["word_svd_22", "0.0030", "получение заказа, возврат, описание и рекомендации"],
            ["char_svd_4", "0.0028", "продавец, доставка, качество и размер"],
            ["word_svd_6", "0.0028", "качество ткани, синтетика, нитки, швы"],
            ["exclaims_per_word", "0.0028", "эмоциональная насыщенность коротких отзывов"],
        ],
    )
    add_paragraph(
        doc,
        "Встроенная важность и permutation importance не обязаны совпадать. Встроенная важность может высоко оценивать "
        "признак, который часто используется в деревьях, но не всегда критичен для итоговой macro-F1. Permutation "
        "importance строже: она проверяет, что произойдет с качеством, если связь между признаком и целевой переменной "
        "разрушить. Поэтому особое значение имеют признаки, которые повторяются в нескольких списках: word_svd_3, "
        "pos_neg_top_diff, num_negations, char_svd_8 и word_svd_6. Случайный baseline-признак не дает снижения качества, "
        "что подтверждает, что верхние позиции не являются техническим шумом."
    )

    add_page_break(doc)
    add_paragraph(doc, "3.3. Расшифровка SVD-компонент", bold=True, first_line_cm=0)
    add_table_caption(doc, "Таблица 7 - Содержательная интерпретация ключевых SVD-компонент")
    add_table(
        doc,
        ["Компонента", "Характерные сигналы", "Осторожная интерпретация"],
        [
            [
                "word_svd_3",
                "«не соответствует», «размер не соответствует», «описанию», «фото», «цвет», «на фото»",
                "ось соответствия ожиданиям: описание, фотография, цвет и размер",
            ],
            [
                "word_svd_6",
                "«качество», «не очень», «плохое», «ужасное», «синтетика», «нитки», «торчат»",
                "качество материала и исполнения",
            ],
            [
                "word_svd_22",
                "«быстро», «получила», «не пришел», «деньги не вернули», «рекомендую»",
                "доставка, получение заказа, возврат денег и рекомендация",
            ],
            [
                "char_svd_1",
                "фрагменты «верну», «вернул», «деньги», а также фрагменты слов о размере и качестве",
                "возврат денег и конфликтные ситуации после покупки",
            ],
            [
                "char_svd_8",
                "фрагменты «достав», «отлич», «спасибо»",
                "противопоставление доставки и позитивной благодарственной лексики",
            ],
            [
                "char_svd_34",
                "фрагменты «супер», «спор», «открыла спор»",
                "полярность между явной похвалой и спорной ситуацией",
            ],
        ],
    )
    add_paragraph(
        doc,
        "SVD-компоненты получились именно такими не потому, что им заранее задали смысл, а потому что в корпусе коротких "
        "отзывов повторяются устойчивые группы слов. Для одежды часто повторяются жалобы на размер, ткань и соответствие "
        "фото, а для сервиса - доставка, спор и возврат денег. TruncatedSVD сжимает TF-IDF-пространство так, что эти "
        "совместные паттерны становятся отдельными латентными координатами. Поэтому интерпретация SVD полезна, но не "
        "является причинным доказательством: ее нужно подтверждать permutation importance, SHAP и содержательным чтением "
        "примеров отзывов."
    )

    add_paragraph(doc, "3.4. PDP/ICE: реакция модели на изменение признаков", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "PDP/ICE-графики показывают, как меняются вероятности классов при движении по выбранному признаку. Это особенно "
        "важно для бизнес-интерпретации: можно увидеть не только факт важности признака, но и направление эффекта. "
        "Для ручных признаков трактовка ближе к прямой: больше отрицаний или более выраженный дисбаланс негативных слов "
        "повышает вероятность отрицательного класса. Для SVD-компонент трактовка сложнее: движение по оси означает "
        "смещение текста к группе слов и фрагментов, а не добавление одного слова."
    )
    add_image(doc, FIG_PDP_POS_NEG, "Рисунок 4. PDP/ICE для признака pos_neg_top_diff")
    add_image(doc, FIG_PDP_NEG, "Рисунок 5. PDP/ICE для признака num_negations")
    add_image(doc, FIG_PDP_NOT_RECOMMEND, "Рисунок 6. PDP/ICE для признака has_not_recommend")
    add_image(doc, FIG_PDP_NUM_CHARS, "Рисунок 7. PDP/ICE для признака num_chars")
    add_paragraph(
        doc,
        "Для pos_neg_top_diff отрицательные значения соответствуют преобладанию негативной лексики и повышают вероятность "
        "отрицательного класса. Для num_negations рост числа отрицаний также связан с усилением негативного сигнала, "
        "но эффект не является строго линейным: после нескольких отрицаний модель часто уже уверена в негативной оценке. "
        "Признак has_not_recommend полезен как явный бизнес-сигнал: фраза «не рекомендую» не просто ухудшает тональность, "
        "а указывает на риск отказа от повторной покупки. Длина текста работает иначе: длинные отзывы чаще содержат "
        "подробности о доставке, размере или качестве, поэтому она помогает модели отличать информативные жалобы от "
        "коротких эмоциональных реакций."
    )
    add_image(doc, FIG_PDP_WORD_SVD_3, "Рисунок 8. PDP/ICE для латентной компоненты word_svd_3")
    add_image(doc, FIG_PDP_WORD_SVD_6, "Рисунок 9. PDP/ICE для латентной компоненты word_svd_6")
    add_image(doc, FIG_PDP_CHAR_SVD_8, "Рисунок 10. PDP/ICE для латентной компоненты char_svd_8")
    add_paragraph(
        doc,
        "Для SVD-компонент PDP/ICE показывает, что модель реагирует не на отдельные слова, а на движение текста вдоль "
        "латентной темы. word_svd_3 связан с несоответствием ожиданиям: размер, фото, цвет и описание. word_svd_6 "
        "отражает качество материала и исполнения: синтетика, плохая ткань, торчащие нитки. char_svd_8 дополняет "
        "словесные признаки за счет фрагментов слов, что важно для русскоязычных отзывов с опечатками, сокращениями "
        "и разными формами одних и тех же слов."
    )

    add_page_break(doc)
    add_paragraph(doc, "3.5. SHAP-анализ по классам", bold=True, first_line_cm=0)
    add_paragraph(
        doc,
        "SHAP-графики дополняют PDP тем, что показывают вклад признаков в конкретные предсказания и отдельно для каждого "
        "класса. В отрицательном классе наиболее заметны pos_neg_top_diff, word_svd_3, char_svd_1 и num_negations: "
        "модель связывает негатив не только с общим лексическим балансом, но и с конкретными клиентскими проблемами. "
        "В положительном классе важны признаки соответствия ожиданиям, благодарственная лексика, быстрая доставка и "
        "отсутствие жалоб. Нейтральный класс остается наиболее сложным: его SHAP-профиль состоит из смешанных описательных "
        "сигналов, где текст может быть информативным, но не иметь сильной эмоциональной полярности."
    )
    add_image(doc, FIG_SHAP_NEG, "Рисунок 11. SHAP beeswarm для отрицательного класса", width_cm=14.0)
    add_image(doc, FIG_SHAP_NEUTRAL, "Рисунок 12. SHAP beeswarm для нейтрального класса", width_cm=14.0)
    add_image(doc, FIG_SHAP_POS, "Рисунок 13. SHAP beeswarm для положительного класса", width_cm=14.0)

    add_heading(doc, "4. Обсуждение")
    add_paragraph(
        doc,
        "Сравнение с линейной baseline-моделью делает вывод статьи более честным: бустинг не является "
        "единственным способом классификации таких отзывов, но после refit на полном train лучший бустинг XGBoost "
        "слегка превысил baseline по macro-F1. При сильных TF-IDF/SVD-признаках линейная модель все равно остается "
        "очень конкурентной. Однако для задачи бизнес-аналитики важны не только проценты macro-F1. LightGBM позволяет "
        "анализировать нелинейные пороги, взаимодействие признаков "
        "и устойчивость сигналов несколькими независимыми методами интерпретации."
    )
    add_paragraph(
        doc,
        "Главный прикладной результат состоит в том, что модель выделяет не абстрактную «позитивность» или «негативность», "
        "а повторяющиеся группы клиентских проблем. Первая группа - несоответствие ожиданиям: размер, цвет, фото и описание. "
        "Вторая - качество товара: ткань, синтетика, нитки, швы и общая оценка материала. Третья - логистика: доставка, "
        "неполучение заказа и длительное ожидание. Четвертая - постпокупочный конфликт: спор, возврат денег и явный отказ "
        "рекомендовать товар. Именно эти группы полезны для бизнеса: они могут быть превращены в отчеты для категорийного "
        "менеджмента, контроля поставщиков и клиентского сервиса."
    )
    add_paragraph(
        doc,
        "Ограничения исследования также существенны. Во-первых, данные относятся только к одной товарной категории, "
        "поэтому выводы нельзя напрямую переносить на электронику, продукты питания или услуги. Во-вторых, разница между "
        "бустинговыми моделями мала и bootstrap-интервалы пересекаются, поэтому XGBoost корректно называть лучшим только "
        "по точечной оценке после refit, а не статистически доказанным лидером. В-третьих, SVD-компоненты требуют "
        "осторожной интерпретации: они являются латентными осями, а не понятными ручными признаками. В-четвертых, в работе "
        "не проверялись трансформерные модели типа RuBERT, поэтому статья не претендует на сравнение с современными "
        "нейросетевыми решениями по максимальному качеству."
    )

    add_heading(doc, "5. Заключение")
    add_paragraph(
        doc,
        "В работе показано, что анализ тональности отзывов можно использовать как инструмент объяснимой бизнес-аналитики, "
        "а не только как задачу автоматической разметки текста. На корпусе русскоязычных отзывов о женской одежде и "
        "аксессуарах построено гибридное признаковое пространство, обучены линейный baseline и три модели градиентного "
        "бустинга. После refit на полном train лучшую точечную метрику среди бустингов показал XGBoost, а близкая по "
        "качеству LightGBM-модель подробно интерпретирована через feature importance, permutation importance, SHAP и PDP/ICE."
    )
    add_paragraph(
        doc,
        "Подробный анализ признаков показывает, что для клиентов наиболее важны не только общие эмоциональные слова, "
        "но и конкретные товарные и сервисные аспекты: соответствие описанию и фотографии, размер, качество ткани и швов, "
        "доставка, возврат денег и явное нежелание рекомендовать товар. Именно в этом состоит практическая ценность "
        "бустингового подхода: он дает достаточно сильную классификацию и одновременно позволяет объяснить, какие "
        "клиентские pain points стоят за решениями модели. Дальнейшая работа должна включать ручную валидацию выделенных "
        "аспектов, проверку на других категориях и сравнение с трансформерными моделями."
    )

    add_heading(doc, "Список литературы", page_break_before=True)
    refs = [
        "1. Friedman J. H. Greedy Function Approximation: A Gradient Boosting Machine // Annals of Statistics. 2001. Vol. 29. No. 5. P. 1189-1232.",
        "2. Lundberg S. M., Lee S.-I. A Unified Approach to Interpreting Model Predictions // Advances in Neural Information Processing Systems. 2017. Vol. 30. P. 4765-4774.",
        "3. Ke G., Meng Q., Finley T. et al. LightGBM: A Highly Efficient Gradient Boosting Decision Tree // Advances in Neural Information Processing Systems. 2017. Vol. 30.",
        "4. Smetanin S., Komarov M. Sentiment Analysis of Product Reviews in Russian using Convolutional Neural Networks // 2019 IEEE 21st Conference on Business Informatics (CBI). 2019. P. 482-486. DOI: 10.1109/CBI.2019.00062.",
        "5. RuReviews: An Automatically Annotated Sentiment Analysis Dataset for Product Reviews in Russian [Электронный ресурс]. URL: https://github.com/sismetanin/rureviews.",
        "6. Двойникова А. А., Карпов А. А. Аналитический обзор подходов к распознаванию тональности русскоязычных текстовых данных // Информационно-управляющие системы. 2020. № 4. С. 20-30. DOI: 10.31799/1684-8853-2020-4-20-30.",
        "7. Самигулин Т. Р., Джурабаев А. Э. У. Анализ тональности текста методами машинного обучения // Научный результат. Информационные технологии. 2021. Т. 6. № 1. С. 55-62. DOI: 10.18413/2518-1092-2021-6-1-0-7.",
        "8. Mokgwatjane K., Paepae T. An explainable ensemble machine learning approach for multi-domain, multiclass sentiment analysis in Amazon product reviews // Machine Learning with Applications. 2026. Vol. 23. Art. 100825. DOI: 10.1016/j.mlwa.2025.100825.",
        "9. Davoodi L., Mezei J., Heikkilä M. Aspect-based sentiment classification of user reviews to understand customer satisfaction of e-commerce platforms // Electronic Commerce Research. 2026. Vol. 26. P. 1417-1459. DOI: 10.1007/s10660-025-09948-4.",
    ]
    for item in refs:
        add_paragraph(doc, item, first_line_cm=0)

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number(p)
    return doc


def write_title_note() -> None:
    text = f"""# Что дозаполнить в титульнике

Я уже заполнил:

- ФИО: **{STUDENT}**
- Группа: **{GROUP}**
- Курс: **{COURSE}**
- Тема задания: **{TOPIC}**
- Сроки практики: **{PRACTICE_DATES}**
- Предполагаемое место практики: **{PRACTICE_PLACE}**

Проверь и при необходимости исправь:

1. **Место прохождения практики**. Если практика проходила не внутри НГУ, а в другой организации или лаборатории, замени строку на точное официальное название, подразделение и адрес.
2. **Руководитель практики**. Нужны Ф.И.О., ученая степень и звание.
3. **Оценка по итогам защиты отчета**. Обычно это поле заполняется после защиты.
4. **Подпись руководителя** и, если требуется, дата защиты.
5. **Наименование кафедры**. В шаблоне указана кафедра Интеллектуальных систем теплофизики ИИР; если у вас в документах используется другая формулировка, ее лучше привести к официальной.

Что уже подготовлено:

- `{REPORT_NAME}` — готовый отчет по учебной практике.
- `{ARTICLE_NAME}` — готовая научная статья с акцентом на PDP и SHAP.

Реалистичные журналы, под стиль которых статья подходит лучше всего:

- **IEEE Access** — прикладной инженерный и ML-формат, SJR 2024: Q1.
- **Scientific Reports** — сильный междисциплинарный формат, SJR 2024: Q1.
- **Applied Sciences (Switzerland)** — прикладной формат для систем и AI, SJR 2024: Q2.

Если понадобится, следующий шаг — адаптация статьи под конкретные требования выбранного журнала: объем, формат ссылок, требования к аннотации, рисункам и сведениям об авторах.
"""
    (OUT_DIR / TITLE_NOTE_NAME).write_text(text, encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    report_doc = build_report()
    article_doc = build_article()
    report_doc.save(OUT_DIR / REPORT_NAME)
    article_doc.save(OUT_DIR / ARTICLE_NAME)
    write_title_note()
    print((OUT_DIR / REPORT_NAME).resolve())
    print((OUT_DIR / ARTICLE_NAME).resolve())
    print((OUT_DIR / TITLE_NOTE_NAME).resolve())


if __name__ == "__main__":
    main()
